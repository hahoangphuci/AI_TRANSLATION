import os
import re
import unicodedata
import time
import PyPDF2
import docx
import openpyxl
# fpdf is optional; if missing we fallback to text output for PDFs
try:
    from fpdf import FPDF
    HAS_FPDF = True
except Exception:
    FPDF = None
    HAS_FPDF = False
from werkzeug.utils import secure_filename


class ProviderRateLimitError(Exception):
    """Raised when the upstream AI provider indicates a hard rate limit (429 or insufficient credits)."""
    pass

class FileService:
    def __init__(self, translator=None):
        """translator: callable(text, source_lang, target_lang) -> translated_text"""
        from concurrent.futures import ThreadPoolExecutor

        self.upload_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'uploads')
        self.download_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'downloads')
        os.makedirs(self.upload_folder, exist_ok=True)
        os.makedirs(self.download_folder, exist_ok=True)
        self.translator = translator
        # Performance tuning
        try:
            from app import config as app_config
            self.concurrency = getattr(app_config.Config, 'TRANSLATION_CONCURRENCY', 4)
            self.retries = getattr(app_config.Config, 'TRANSLATION_RETRIES', 3)
            self.backoff = getattr(app_config.Config, 'TRANSLATION_BACKOFF', 1.5)
        except Exception:
            self.concurrency = int(os.getenv('TRANSLATION_CONCURRENCY', '4'))
            self.retries = int(os.getenv('TRANSLATION_RETRIES', '3'))
            self.backoff = float(os.getenv('TRANSLATION_BACKOFF', '1.5'))
        self._executor_cls = ThreadPoolExecutor

    def _translate_with_retry(self, text, target_lang):
        """Translate a piece of text with retry/backoff on transient errors.

        IMPORTANT: If a provider rate-limit or "insufficient credits" error is encountered,
        fail fast by raising ProviderRateLimitError so the calling job can abort immediately
        instead of continuing and wasting quota/retries.
        """
        if not self.translator:
            raise RuntimeError('Translator not configured')
        last = None
        attempt = 0
        max_attempts = max(1, self.retries)
        while attempt < max_attempts:
            try:
                out = self.translator(text, 'auto', target_lang)
                return out
            except Exception as e:
                last = e
                err = str(e).lower()
                # Fail fast for provider rate limits or insufficient credits
                if any(k in err for k in ('rate', 'insufficient', 'free-models', '402', 'credit', 'requires more credits')):
                    try:
                        print(f"Provider rate limit or insufficient credits encountered: {e}")
                    except UnicodeEncodeError:
                        print("Provider rate limit encountered: ", repr(e))
                    raise ProviderRateLimitError(str(e))
                # Retry on transient network errors
                if any(k in err for k in ('temporarily', 'timed out', 'timeout', 'connection')):
                    sleep_time = (self.backoff ** attempt)
                    print(f"Translate retry {attempt+1}/{self.retries} after {sleep_time}s due to: {e}")
                    time.sleep(sleep_time)
                    attempt += 1
                    continue
                # Non-retryable errors: break
                break
        # Final attempt to raise helpful error
        raise last

    
    def process_document(self, file_path, target_lang, progress_callback=None):
        filename = os.path.basename(file_path)
        name, ext = os.path.splitext(filename)
        
        if ext.lower() == '.pdf':
            return self._process_pdf(file_path, target_lang, progress_callback)
        elif ext.lower() == '.docx':
            return self._process_docx(file_path, target_lang, progress_callback)
        elif ext.lower() == '.xlsx':
            return self._process_xlsx(file_path, target_lang, progress_callback)
        elif ext.lower() == '.txt':
            return self._process_txt(file_path, target_lang, progress_callback)
        else:
            raise ValueError("Unsupported file type")
    
    def _process_pdf(self, file_path, target_lang, progress_callback=None):
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            total_pages = len(pdf_reader.pages)
            for i, page in enumerate(pdf_reader.pages):
                text += page.extract_text() + "\n"
                if progress_callback:
                    progress_callback(int(5 + (i / max(1, total_pages)) * 20), f"Extracting page {i+1}/{total_pages}")
        
            if progress_callback:
                progress_callback(25, "Translating PDF text...")
            # Split by pages (already done) and translate pages in parallel
            pages = text.split('\f') if '\f' in text else text.split('\n\f') if '\n\f' in text else text.split('\n\n')
            with self._executor_cls(max_workers=self.concurrency) as ex:
                futures = [ex.submit(self._translate_with_retry, p, target_lang) for p in pages]
                translated_pages = []
                for fut in futures:
                    try:
                        translated_pages.append(fut.result())
                    except ProviderRateLimitError:
                        print("Provider rate limit detected during PDF page translation, aborting job.")
                        raise
                    except Exception as e:
                        print(f"PDF page translation failed: {e}")
                        translated_pages.append('')
            translated_text = '\n\n'.join(translated_pages)

        # If FPDF is available, build a PDF; otherwise fallback to a TXT file (no layout preservation)
        if HAS_FPDF and FPDF:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Split text into lines to fit page
            lines = translated_text.split('\n')
            for line in lines:
                # Handle long lines
                while len(line) > 0:
                    if pdf.get_string_width(line) < 180:  # Approximate page width
                        pdf.cell(0, 10, txt=line, ln=True)
                        break
                    else:
                        # Find a good break point
                        words = line.split()
                        current_line = ""
                        for word in words:
                            if pdf.get_string_width(current_line + " " + word) < 180:
                                current_line += " " + word if current_line else word
                            else:
                                pdf.cell(0, 10, txt=current_line, ln=True)
                                current_line = word
                        if current_line:
                            pdf.cell(0, 10, txt=current_line, ln=True)
                        line = ""

            # Ensure output filename has .pdf extension
            output_filename = f"translated_{os.path.basename(file_path)}"
            if not output_filename.lower().endswith('.pdf'):
                output_filename += '.pdf'
            output_path = os.path.join(self.download_folder, output_filename)
            pdf.output(output_path)
        else:
            # Fallback: save plain text and include a note
            if progress_callback:
                progress_callback(60, "PDF library not available, writing plain text fallback")
            output_filename = f"translated_{os.path.basename(file_path)}"
            if not output_filename.lower().endswith('.txt'):
                output_filename += '.txt'
            output_path = os.path.join(self.download_folder, output_filename)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("NOTE: PDF rebuild not available on this system. Install 'fpdf' to get translated PDF output.\n\n")
                f.write(translated_text)

        if progress_callback:
            progress_callback(100, "Completed")
        return output_path
    
    def _process_docx(self, file_path, target_lang, progress_callback=None):
        # Modify original document in-place so styles/images/relationships are preserved
        doc = docx.Document(file_path)

        def distribute_text_to_runs(translated: str, runs_texts):
            """
            Heuristic: keep the same run count/styles by distributing the translated
            paragraph text back into existing runs proportionally by original run length.
            """
            if translated is None:
                translated = ""
            translated = str(translated)

            # Only consider runs that had some text (including whitespace) for distribution
            lengths = [len(t) for t in runs_texts]
            total = sum(lengths)
            if total <= 0:
                # If all runs are empty, put everything into the first run
                return [translated] + [""] * (len(runs_texts) - 1)

            # Initial proportional allocation
            alloc = []
            used = 0
            for i, ln in enumerate(lengths):
                if i == len(lengths) - 1:
                    take = len(translated) - used
                else:
                    take = int(round((ln / total) * len(translated)))
                    take = max(0, min(take, len(translated) - used))
                alloc.append(translated[used : used + take])
                used += take

            # Fix rounding drift
            if used < len(translated):
                alloc[-1] += translated[used:]
            elif used > len(translated):
                # Trim from the end if we overshot
                overflow = used - len(translated)
                if overflow > 0 and alloc[-1]:
                    alloc[-1] = alloc[-1][:-overflow]

            # Ensure same length
            if len(alloc) < len(runs_texts):
                alloc.extend([""] * (len(runs_texts) - len(alloc)))
            return alloc[: len(runs_texts)]

        def translate_paragraph_runs(paragraph, idx=None, total=None):
            # Translate at paragraph-level for better quality, then map back to runs to keep formatting.
            runs = list(paragraph.runs)
            if not runs:
                return

            original_texts = [r.text or "" for r in runs]
            paragraph_text = "".join(original_texts)
            if not paragraph_text.strip():
                return

            try:
                translated_para = self._translate_text(paragraph_text, target_lang)
            except ProviderRateLimitError:
                # Critical: if provider is rate-limited stop the entire document job
                print("Provider rate limit detected during paragraph translation, raising to abort job.")
                raise
            except Exception as e:
                print(f"Translator failed for paragraph: {e}")
                translated_para = paragraph_text

            pieces = distribute_text_to_runs(translated_para, original_texts)
            for run, piece in zip(runs, pieces):
                run.text = piece
            if progress_callback and idx is not None and total is not None:
                progress_callback(10 + int((idx / total) * 70), f"Translating paragraph {idx+1}/{total}")

        # Body paragraphs
        paragraphs = [p for p in doc.paragraphs]
        total = len(paragraphs) if paragraphs else 1
        # Translate in parallel using ThreadPoolExecutor
        with self._executor_cls(max_workers=self.concurrency) as ex:
            futures = {}
            for idx, para in enumerate(paragraphs):
                # capture original run texts
                runs_texts = [r.text or "" for r in para.runs]
                paragraph_text = "".join(runs_texts)
                if not paragraph_text.strip():
                    continue
                futures[ex.submit(self._translate_with_retry, paragraph_text, target_lang)] = (idx, para, runs_texts)

            completed = 0
            for fut in list(futures):
                try:
                    res = fut.result()
                    idx, para, runs_texts = futures[fut]
                    pieces = distribute_text_to_runs(res, runs_texts)
                    for run, piece in zip(list(para.runs), pieces):
                        run.text = piece
                except ProviderRateLimitError as e:
                    print("Provider rate limit detected during paragraph processing, aborting job.")
                    raise
                except Exception as e:
                    print(f"Paragraph translation failed: {e}")
                completed += 1
                if progress_callback:
                    progress_callback(10 + int((completed / max(1, total)) * 70), f"Translating paragraph {completed}/{total}")

        # Tables: translate cell-by-cell, preserve cell formatting
        for table in doc.tables:
            for r in range(len(table.rows)):
                for c in range(len(table.columns)):
                    cell = table.rows[r].cells[c]
                    for p_idx, p in enumerate(cell.paragraphs):
                        translate_paragraph_runs(p, p_idx, len(cell.paragraphs))

        # Headers and footers
        try:
            for section in doc.sections:
                header = section.header
                for p_idx, p in enumerate(header.paragraphs):
                    translate_paragraph_runs(p, p_idx, len(header.paragraphs))
                footer = section.footer
                for p_idx, p in enumerate(footer.paragraphs):
                    translate_paragraph_runs(p, p_idx, len(footer.paragraphs))
        except Exception:
            # ignore headers/footers issues
            pass

        # Ensure output filename has .docx extension
        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'
        output_path = os.path.join(self.download_folder, output_filename)

        # Save and validate
        doc.save(output_path)


        # Validate produced DOCX — if invalid, write a plain text fallback to avoid corrupt file being returned
        try:
            # Try opening the saved file with python-docx to validate
            docx.Document(output_path)
        except Exception as e:
            # Create a text fallback containing translated paragraphs and table text
            if progress_callback:
                progress_callback(95, "DOCX validation failed, writing fallback text file")
            fallback_filename = output_filename
            if not fallback_filename.lower().endswith('.txt'):
                fallback_filename += '.txt'
            fallback_path = os.path.join(self.download_folder, fallback_filename)
            # Collect text from doc
            lines = []
            for p in doc.paragraphs:
                lines.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        lines.append(cell.text)
            with open(fallback_path, 'w', encoding='utf-8') as f:
                f.write("NOTE: DOCX creation failed on server. Showing plain text fallback below.\n\n")
                f.write('\n'.join(lines))
            output_path = fallback_path

        if progress_callback:
            progress_callback(100, "Completed")
        return output_path
    
    def _process_xlsx(self, file_path, target_lang, progress_callback=None):
        # Translate in-place to preserve styles, merged cells, formulas, column widths, etc.
        wb = openpyxl.load_workbook(file_path)

        # Count total cells (rough progress)
        total_cells = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for _ in ws.iter_rows():
                total_cells += 1
        total_cells = total_cells or 1

        # Collect cells to translate
        to_translate = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    try:
                        is_formula = (cell.data_type == 'f') or (
                            isinstance(cell.value, str) and cell.value.startswith("=")
                        )
                    except Exception:
                        is_formula = False

                    if (not is_formula) and isinstance(cell.value, str) and cell.value.strip():
                        to_translate.append(cell)

        total = len(to_translate) or 1
        processed = 0
        # Translate cells in parallel
        with self._executor_cls(max_workers=self.concurrency) as ex:
            futures = {ex.submit(self._translate_with_retry, cell.value, target_lang): cell for cell in to_translate}
            for fut in futures:
                try:
                    translated = fut.result()
                    cell = futures[fut]
                    cell.value = translated
                except ProviderRateLimitError:
                    print("Provider rate limit detected during cell translation, aborting job.")
                    raise
                except Exception as e:
                    print(f"Cell translation failed: {e}")
                processed += 1
                if progress_callback:
                    progress_callback(10 + int((processed / total) * 80), f"Translating cells {processed}/{total}")

        # Ensure output filename has .xlsx extension
        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.xlsx'):
            output_filename += '.xlsx'
        output_path = os.path.join(self.download_folder, output_filename)
        wb.save(output_path)
        if progress_callback:
            progress_callback(100, "Completed")
        return output_path
    
    def _process_txt(self, file_path, target_lang, progress_callback=None):
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        if progress_callback:
            progress_callback(25, "Translating text file...")

        # Split into paragraphs then chunk long paragraphs to avoid provider length limits
        paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        max_chars = 3000
        chunks = []
        for p in paras:
            if len(p) <= max_chars:
                chunks.append(p)
            else:
                parts = re.split(r'(?<=[.!?])\s+', p)
                cur = ''
                for part in parts:
                    if len(cur) + len(part) + 1 <= max_chars:
                        cur = (cur + ' ' + part).strip() if cur else part
                    else:
                        if cur:
                            chunks.append(cur)
                        cur = part
                if cur:
                    # If still too long, slice it
                    while len(cur) > max_chars:
                        chunks.append(cur[:max_chars])
                        cur = cur[max_chars:]
                    if cur:
                        chunks.append(cur)

        # Translate chunks in parallel
        translated_parts = []
        with self._executor_cls(max_workers=self.concurrency) as ex:
            futures = [ex.submit(self._translate_with_retry, c, target_lang) for c in chunks]
            total = len(futures) or 1
            for i, fut in enumerate(futures, start=1):
                try:
                    res = fut.result()
                    translated_parts.append(res)
                except ProviderRateLimitError as e:
                    try:
                        print("Provider rate limit detected during text file translation, aborting job.")
                    except UnicodeEncodeError:
                        print("Provider rate limit detected during text file translation, aborting job.")
                    raise
                except Exception as e:
                    print(f"Chunk translation failed: {e}")
                    translated_parts.append('')
                if progress_callback:
                    progress_callback(25 + int((i / total) * 70), f"Translating chunk {i}/{total}")

        translated_text = '\n\n'.join(translated_parts)

        output_filename = f"translated_{os.path.basename(file_path)}"
        if not output_filename.lower().endswith('.txt'):
            output_filename += '.txt'
        output_path = os.path.join(self.download_folder, output_filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_text)
        if progress_callback:
            progress_callback(100, "Completed")
        return output_path
    
    def _sanitize_text(self, text: str) -> str:
        if not isinstance(text, str):
            try:
                text = str(text)
            except Exception:
                return ''
        # Normalize unicode and remove control characters that break XML/docx
        text = unicodedata.normalize('NFC', text)
        # Remove C0 control characters except tab/newline/carriage return
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
        # Collapse weird zero-width/formatting if any
        text = re.sub(r'[\u200B-\u200F\u2028\u2029]', ' ', text)
        return text

    def _translate_text(self, text, target_lang):
        # Use injected translator with retry/backoff
        if self.translator:
            out = self._translate_with_retry(text, target_lang)
            return self._sanitize_text(out)
        raise RuntimeError("Translator is not configured")