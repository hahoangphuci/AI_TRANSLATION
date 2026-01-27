"""
Translate first N paragraphs from uploads/phuong_orig.docx and save a partial translated docx to downloads/translated_phuong_sample.docx
Run:
& "C:/Users/HA HOANG PHUC/AppData/Local/Programs/Python/Python313/python.exe" scripts/translate_phuong_sample.py
"""
from pathlib import Path
from app.services.translation_service import TranslationService
from docx import Document

N = 10
src = Path('uploads/phuong_orig.docx')
if not src.exists():
    print('Source file not found:', src)
    raise SystemExit(1)

svc = TranslationService()
print('TranslationService initialized')

doc = Document(src)
print('Total paragraphs in doc:', len(doc.paragraphs))

# Translate first N paragraphs (run-level) and write to a new document preserving structure
out_doc = Document(src)  # start from original to preserve styles
for i in range(min(N, len(doc.paragraphs))):
    p = out_doc.paragraphs[i]
    orig_p = doc.paragraphs[i]
    print(f'Paragraph {i+1} original preview:', orig_p.text[:120])
    for run_idx, run in enumerate(orig_p.runs):
        txt = run.text
        if not txt.strip():
            continue
        try:
            translated = svc.translate_text(txt, 'vi', 'en')
        except Exception as e:
            print('Translation failed for run:', e)
            translated = txt
        # sanitize printed preview
        print('  run', run_idx, '->', translated[:120])
        # replace corresponding run in out_doc
        try:
            out_doc.paragraphs[i].runs[run_idx].text = translated
        except Exception:
            # fallback: replace whole paragraph text
            out_doc.paragraphs[i].text = out_doc.paragraphs[i].text.replace(txt, translated)

# Save partial translated document
out_path = Path('downloads/translated_phuong_sample.docx')
out_doc.save(out_path)
print('Saved sample translated file:', out_path)
print('Done')
